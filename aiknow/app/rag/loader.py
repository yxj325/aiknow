# -*- coding: utf-8 -*-
"""
文档加载器 - 支持 PDF / Word / Excel / Markdown / 纯文本 / HTML
"""

import os
from pathlib import Path
from typing import List
import logging

logger = logging.getLogger(__name__)


class Document:
    def __init__(self, content: str, metadata: dict = None):
        self.content = content
        self.metadata = metadata or {}

    def __repr__(self):
        return f"Document(meta={self.metadata}, len={len(self.content)})"


class DocumentLoader:
    SUPPORTED_EXT = {".pdf", ".docx", ".doc", ".xlsx", ".xls",
                     ".md", ".markdown", ".txt", ".html", ".htm"}

    def load(self, path: str) -> List[Document]:
        path = str(path)
        ext = Path(path).suffix.lower()
        if ext not in self.SUPPORTED_EXT:
            logger.warning("不支持的文件类型: %s", ext)
            return []
        loader_map = {
            ".pdf": self._load_pdf, ".docx": self._load_docx, ".doc": self._load_docx,
            ".xlsx": self._load_excel, ".xls": self._load_excel,
            ".md": self._load_markdown, ".markdown": self._load_markdown,
            ".txt": self._load_text, ".html": self._load_html, ".htm": self._load_html,
        }
        loader = loader_map.get(ext)
        if not loader:
            return []
        try:
            docs = loader(path)
            for d in docs:
                d.metadata.setdefault("source", os.path.basename(path))
                d.metadata.setdefault("filepath", path)
                d.metadata.setdefault("filetype", ext)
            logger.info("加载 %s: %d 文档块", path, len(docs))
            return docs
        except Exception as e:
            logger.error("加载失败 %s: %s", path, e)
            return []

    def load_batch(self, paths: List[str]) -> List[Document]:
        docs = []
        for p in paths:
            docs.extend(self.load(p))
        return docs

    def load_dir(self, directory: str, recursive: bool = True) -> List[Document]:
        pattern = "**/*" if recursive else "*"
        paths = []
        for ext in self.SUPPORTED_EXT:
            paths.extend(str(p) for p in Path(directory).glob(f"{pattern}{ext}") if p.is_file())
        return self.load_batch(paths)

    def _load_pdf(self, path: str) -> List[Document]:
        import fitz
        docs = []
        doc = fitz.open(path)
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                docs.append(Document(text, {"page": page_num + 1, "page_label": page.get_label() or str(page_num + 1)}))
        doc.close()
        return docs

    def _load_docx(self, path: str) -> List[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        return [Document(text, {"type": "document"})] if text else []

    def _load_excel(self, path: str) -> List[Document]:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        docs = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                docs.append(Document("\n".join(rows), {"sheet": sheet_name, "type": "table"}))
        wb.close()
        return docs

    def _load_markdown(self, path: str) -> List[Document]:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return [Document(text, {"type": "markdown"})]

    def _load_text(self, path: str) -> List[Document]:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return [Document(text, {"type": "text"})]

    def _load_html(self, path: str) -> List[Document]:
        from html.parser import HTMLParser
        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.texts = []
                self._skip = False
            def handle_starttag(self, tag, attrs):
                if tag in ("script", "style"):
                    self._skip = True
            def handle_endtag(self, tag):
                if tag in ("script", "style"):
                    self._skip = False
            def handle_data(self, data):
                if not self._skip and data.strip():
                    self.texts.append(data.strip())
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        extractor = TextExtractor()
        extractor.feed(content)
        text = "\n".join(extractor.texts)
        return [Document(text, {"type": "html"})] if text else []


def load_document(path: str) -> List[Document]:
    return DocumentLoader().load(path)

def load_directory(directory: str, recursive: bool = True) -> List[Document]:
    return DocumentLoader().load_dir(directory, recursive=recursive)

def load_text(text: str, metadata: dict = None) -> List[Document]:
    return [Document(text, metadata or {})]
